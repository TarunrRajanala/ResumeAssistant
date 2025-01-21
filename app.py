from flask import Flask, request, jsonify, render_template, send_from_directory
from flask_cors import CORS
from groq import Groq
import os
from docx import Document
from werkzeug.utils import secure_filename
import magic  # for MIME type checking
import logging
from logging.handlers import RotatingFileHandler
from dotenv import load_dotenv
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docxtpl import DocxTemplate
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Load environment variables from .env file
load_dotenv()

# Define constants first
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'
MAX_CONTENT_LENGTH = 32 * 1024 * 1024  # 32MB limit
ALLOWED_EXTENSIONS = {'docx'}

# Create required directories
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# Initialize Flask app
app = Flask(__name__)

# Update CORS configuration for production
if os.environ.get('FLASK_ENV') == 'production':
    CORS(app, resources={r"/*": {"origins": ["*"]}})  # Update with your domain later
    UPLOAD_FOLDER = '/var/app/current/uploads'
    OUTPUT_FOLDER = '/var/app/current/outputs'
else:
    CORS(app)

# Add configurations
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH
app.config['STATIC_FOLDER'] = 'static'
app.config['TEMPLATES_FOLDER'] = 'templates'
app.config['PREFERRED_URL_SCHEME'] = 'https'
app.config['SESSION_COOKIE_SECURE'] = True
app.config['SESSION_COOKIE_HTTPONLY'] = True

# Initialize Groq client
client = Groq(api_key=os.environ.get("GROQ_API_KEY"))

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def read_docx(file_path):
    try:
        doc = Document(file_path)
        return "\n".join(paragraph.text for paragraph in doc.paragraphs)
    except Exception as e:
        print(f"Error reading .docx file: {e}")
        exit(1)

def write_docx(content, file_path):
    """Write content to a .docx file"""
    try:
        doc = Document()
        for paragraph in content.split("\n"):
            if paragraph.strip():  # Only add non-empty paragraphs
                doc.add_paragraph(paragraph.strip())
        doc.save(file_path)
        return True
    except Exception as e:
        app.logger.error(f"Error writing .docx file: {e}")
        return False

def write_pdf(content, file_path):
    """Write content to a PDF file"""
    try:
        # Create the PDF document
        doc = SimpleDocTemplate(
            file_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )

        # Create styles
        styles = getSampleStyleSheet()
        normal_style = styles['Normal']
        
        # Create story (content elements)
        story = []
        
        # Split content into paragraphs and add to story
        for paragraph in content.split('\n'):
            if paragraph.strip():
                p = Paragraph(paragraph.strip(), normal_style)
                story.append(p)
                story.append(Spacer(1, 12))  # Add space between paragraphs
        
        # Build the PDF
        doc.build(story)
        return True
    except Exception as e:
        app.logger.error(f"Error writing PDF file: {e}")
        return False

#AI Cover letter generation
def generate_cover_letter(user_name, job_title, company_name, job_description, user_resume):
    prompt = (
        f"Dear Hiring Manager,\n\n"
        f"Write a professional and tailored cover letter for {user_name}, applying for the position of {job_title} "
        f"at {company_name}. Use the following job description and resume for context:\n\n"
        f"Job Description:\n{job_description}\n\n"
        f"Resume:\n{user_resume}\n\n"
        f"Highlight specific skills, achievements, and experiences that align with the job requirements. "
        f"Make the content ready to submit, avoiding placeholders like '[Your Address]' or '[Recipient's Name]' rather try to pull this data from the resume if available. "
        f"Include quantifiable accomplishments and a professional tone."
        f"Keep the cover letter concise and limited to 3–4 paragraphs. Avoid overly verbose sections."
    )
    try:
        completion = client.chat.completions.create(
            model="llama3-70b-8192",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=1024,
            top_p=0.9,
        )
        return completion.choices[0].message.content
    except Exception as e:
        app.logger.error(f"Error generating cover letter: {e}")
        return None

#AI Resume customization
def generate_custom_resume(user_resume, job_description):
    prompt = (
        f"Customize the following resume to align with this job description:\n\n"
        f"Resume should follow the McCombs resume format. customize the content to align with the job description. "
        f"ensure there are proper sections and headers as shown in the McCombs resume format."
        f"Add or edit descriptions as seen fit in order to make the resume more relevant to the job description."
        f"Quantify the accomplishments and experiences in the resume to make it more relevant to the job description."
        f"Trim off irrelevant information from the resume if no value is added."
        f"End goal is to make the resume more relevant to the job description in order to increase the chance of getting to the next round of interviews."
        f"Customized Resume should only be 1 page long at max."
        f"Job Description:\n{job_description}\n\n"
        f"Resume:\n{user_resume}"
    )
    try:
        completion = client.chat.completions.create(
            model="llama3-70b-8192",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=1024,
            top_p=0.9,
        )
        return completion.choices[0].message.content
    except Exception as e:
        app.logger.error(f"Error generating customized resume: {e}")
        return None

def generate_resume_prompt(job_description, resume):
    return f"""You are a professional resume writer. Generate resume content that will fit EXACTLY into a McCombs template with these strict sections and formatting, ensure the horizontal sectional bars remain as well:

    Output Format Requirements:
    1. First line: Full Name only
    2. Second line: Address | Phone | Email | LinkedIn (separated by pipes)

    3. EDUCATION
    _____________________________________________________________
    SchoolName                                                    City, State
    Degree, Major                                                Month Year
    • GPA and Honors (if applicable)
    • Key coursework (2-3 most relevant)

    4. EXPERIENCE
    _____________________________________________________________
    CompanyName                                                   City, State
    Position Title                                               Month Year - Month Year
    • Achievement-focused bullet (start with action verb)
    • 2-3 bullets per position, quantify results (%, $, #)

    5. LEADERSHIP & ACTIVITIES/PROJECTS
    _____________________________________________________________
    Organization/ProjectName                                      City, State
    Role                                                         Month Year - Month Year
    • 1-2 impactful bullets per activity

    6. SKILLS & INTERESTS
    _____________________________________________________________
    • Technical Skills: List relevant technical skills
    • Languages: List language proficiencies
    • Interests: Brief list of professional interests

    IMPORTANT RULES:
    - Each section header must be in ALL CAPS followed by a horizontal line
    - Content MUST fit on one page
    - Use ONLY the sections and format specified above
    - Align locations and dates to the right margin
    - Use bullet points (•) for all lists
    - Tailor content to match job requirements
    - Use strong action verbs
    - Quantify achievements where possible
    - Maintain consistent spacing between sections

    Current Resume:
    {resume}

    Job Description:
    {job_description}

    Generate the resume content following the EXACT format above, including section headers and horizontal lines."""

def create_cover_letter_template():
    document = Document()
    
    # Set page margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    
    # Define styles
    styles = document.styles
    
    # Header style (for name and contact)
    header_style = styles.add_style('HeaderStyle', WD_STYLE_TYPE.PARAGRAPH)
    header_style.font.name = 'Arial'
    header_style.font.size = Pt(12)
    header_style.paragraph_format.space_after = Pt(0)
    header_style.paragraph_format.line_spacing = 1.0
    
    # Date style
    date_style = styles.add_style('DateStyle', WD_STYLE_TYPE.PARAGRAPH)
    date_style.font.name = 'Arial'
    date_style.font.size = Pt(12)
    date_style.paragraph_format.space_before = Pt(24)
    date_style.paragraph_format.space_after = Pt(12)
    
    # Recipient style
    recipient_style = styles.add_style('RecipientStyle', WD_STYLE_TYPE.PARAGRAPH)
    recipient_style.font.name = 'Arial'
    recipient_style.font.size = Pt(12)
    recipient_style.paragraph_format.space_after = Pt(12)
    
    # Content style
    content_style = styles.add_style('ContentStyle', WD_STYLE_TYPE.PARAGRAPH)
    content_style.font.name = 'Arial'
    content_style.font.size = Pt(12)
    content_style.paragraph_format.space_after = Pt(12)
    content_style.paragraph_format.line_spacing = 1.15
    
    return document

def format_cover_letter(document, content, user_info):
    """
    Format the cover letter content into a professional .docx
    """
    # Add header with name and contact
    header = document.add_paragraph(style='HeaderStyle')
    header.add_run(user_info['name'].upper()).bold = True
    
    # Add contact info
    contact = document.add_paragraph(style='HeaderStyle')
    contact.text = f"{user_info['email']} • {user_info['phone']}"
    if 'github' in user_info:
        contact.text += f" • {user_info['github']}"
    if 'address' in user_info:
        contact.text += f" • {user_info['address']}"
    
    # Add current date
    date = document.add_paragraph(style='DateStyle')
    date.text = datetime.now().strftime("%B %d, %Y")
    
    # Add recipient info
    recipient = document.add_paragraph(style='RecipientStyle')
    recipient.text = f"{user_info['company_name']}\nHiring Manager"
    
    # Add salutation
    salutation = document.add_paragraph(style='ContentStyle')
    salutation.text = "Dear Hiring Manager,"
    
    # Add content paragraphs
    paragraphs = content.split('\n\n')
    for para in paragraphs:
        if para.strip():
            p = document.add_paragraph(style='ContentStyle')
            p.text = para.strip()
    
    # Add closing
    closing = document.add_paragraph(style='ContentStyle')
    closing.text = f"Sincerely,\n{user_info['name']}"

#API Endpoints
@app.route('/generate-cover-letter', methods=['POST'])
def generate_cover_letter_endpoint():
    """
    Endpoint to generate a cover letter. Accepts job description and resume files,
    along with JSON input for user name, job title, and company name.
    """
    try:
        # Check for required files in the request
        if 'job_description' not in request.files or 'resume' not in request.files:
            return jsonify({"error": "Both job description and resume files are required"}), 400

        # Save and read the job description
        job_description_file = request.files['job_description']
        if not allowed_file(job_description_file.filename):
            return jsonify({"error": "Invalid file format for job description"}), 400
        
        job_description_path = os.path.join(app.config['UPLOAD_FOLDER'], 
                                          secure_filename(job_description_file.filename))
        job_description_file.save(job_description_path)
        job_description = read_docx(job_description_path)

        # Save and read the resume
        resume_file = request.files['resume']
        if not allowed_file(resume_file.filename):
            return jsonify({"error": "Invalid file format for resume"}), 400
        
        resume_path = os.path.join(app.config['UPLOAD_FOLDER'], 
                                  secure_filename(resume_file.filename))
        resume_file.save(resume_path)
        user_resume = read_docx(resume_path)

        # Get form data
        user_name = request.form.get("user_name")
        job_title = request.form.get("job_title")
        company_name = request.form.get("company_name")

        if not all([user_name, job_title, company_name]):
            return jsonify({"error": "Missing required parameters"}), 400

        # Generate cover letter
        cover_letter = generate_cover_letter(user_name, job_title, company_name, 
                                          job_description, user_resume)

        if cover_letter:
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            output_filename = f"cover_letter_{timestamp}.docx"
            output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
            
            write_docx(cover_letter, output_path)
            
            # Clean up uploaded files
            cleanup_files(job_description_path, resume_path)
            
            return jsonify({
                "message": "Cover letter generated successfully",
                "file_path": output_filename
            })
        else:
            return jsonify({"error": "Failed to generate cover letter"}), 500

    except Exception as e:
        app.logger.error(f"Error in generate_cover_letter_endpoint: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/customize-resume', methods=['POST'])
def customize_resume():
    try:
        if 'resume' not in request.files:
            return jsonify({'error': 'No resume file uploaded'}), 400
            
        resume_file = request.files['resume']
        if not resume_file.filename.endswith('.docx'):
            return jsonify({'error': 'Please upload a .docx file'}), 400

        # Read the uploaded resume
        input_doc = Document(resume_file)
        
        # Create new document with our formatting
        output_doc = create_resume_template()
        
        # Format the content while maintaining the user's information
        format_resume(input_doc.paragraphs, output_doc)
        
        # Save the formatted document
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], 'formatted_resume.docx')
        output_doc.save(output_path)
        
        return jsonify({'file_path': 'formatted_resume.docx'})

    except Exception as e:
        return jsonify({'error': str(e)}), 500

# Setup logging
def setup_logging():
    if not os.path.exists('logs'):
        os.makedirs('logs')
    
    file_handler = RotatingFileHandler('logs/app.log', maxBytes=10240, backupCount=10)
    file_handler.setFormatter(logging.Formatter(
        '%(asctime)s %(levelname)s: %(message)s [in %(pathname)s:%(lineno)d]'
    ))
    file_handler.setLevel(logging.INFO)
    app.logger.addHandler(file_handler)
    app.logger.setLevel(logging.INFO)
    app.logger.info('Smart Career Assistant startup')

# Add function for cleaning up uploaded files
def cleanup_files(*file_paths):
    """Clean up uploaded files after processing"""
    for file_path in file_paths:
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
        except Exception as e:
            app.logger.error(f"Error cleaning up file {file_path}: {e}")

@app.route('/')
def index():
    try:
        return render_template('index.html')
    except Exception as e:
        app.logger.error(f"Error rendering index: {e}")
        return str(e), 500

@app.route('/download/<path:filename>')
def download_file(filename):
    """Download generated files"""
    try:
        app.logger.info(f"Attempting to download file: {filename}")
        file_path = os.path.join(app.config['OUTPUT_FOLDER'], filename)
        
        if not os.path.exists(file_path):
            app.logger.error(f"File not found: {file_path}")
            return jsonify({"error": f"File not found: {filename}"}), 404
            
        app.logger.info(f"File found, sending: {file_path}")
        return send_from_directory(
            app.config['OUTPUT_FOLDER'],
            filename,
            as_attachment=True
        )
    except Exception as e:
        app.logger.error(f"Error downloading file {filename}: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/test')
def test():
    return jsonify({"status": "Server is running"})

@app.route('/health')
def health_check():
    return jsonify({"status": "healthy"}), 200

# Add after the app initialization
@app.after_request
def after_request(response):
    response.headers.add('Access-Control-Allow-Origin', '*')
    response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
    response.headers.add('Access-Control-Allow-Methods', 'GET,PUT,POST,DELETE,OPTIONS')
    return response

@app.errorhandler(403)
def forbidden_error(error):
    app.logger.error(f"403 error: {error}")
    return jsonify({"error": "Access forbidden"}), 403

@app.errorhandler(500)
def internal_error(error):
    app.logger.error(f'Server Error: {error}')
    return jsonify({"error": "Internal Server Error"}), 500

@app.errorhandler(404)
def not_found_error(error):
    app.logger.error(f'Not Found: {error}')
    return jsonify({"error": "Resource Not Found"}), 404

# Add after app initialization
app.config['SECRET_KEY'] = os.urandom(24)
app.config['TEMPLATES_AUTO_RELOAD'] = True

def ensure_directories():
    """Ensure required directories exist and are writable"""
    for directory in [UPLOAD_FOLDER, OUTPUT_FOLDER]:
        try:
            os.makedirs(directory, exist_ok=True)
            # Test if directory is writable
            test_file = os.path.join(directory, 'test.txt')
            with open(test_file, 'w') as f:
                f.write('test')
            os.remove(test_file)
        except Exception as e:
            app.logger.error(f"Error with directory {directory}: {e}")
            raise

def create_mccombs_template():
    doc = Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    
    # Define styles
    styles = doc.styles
    
    # Name style
    name_style = styles.add_style('Name', WD_STYLE_TYPE.PARAGRAPH)
    name_style.font.name = 'Times New Roman'
    name_style.font.size = Pt(16)
    name_style.font.bold = True
    
    # Header style
    header_style = styles.add_style('Header', WD_STYLE_TYPE.PARAGRAPH)
    header_style.font.name = 'Times New Roman'
    header_style.font.size = Pt(12)
    header_style.font.bold = True
    
    # Normal text style
    normal_style = styles.add_style('Normal Text', WD_STYLE_TYPE.PARAGRAPH)
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(10)
    
    # Add template sections
    doc.add_paragraph('${name}', style='Name')
    doc.add_paragraph('${contact}', style='Normal Text')
    doc.add_paragraph()
    
    # Add each section
    sections = ['EDUCATION', 'EXPERIENCE', 'LEADERSHIP & ACTIVITIES', 'SKILLS & INTERESTS']
    for section in sections:
        p = doc.add_paragraph(section, style='Header')
        p.runs[0].bold = True
        border = p.paragraph_format.border_bottom
        border.width = Pt(1)
        doc.add_paragraph('${' + section.lower().replace(' & ', '_').replace(' ', '_') + '}', 
                         style='Normal Text')
        doc.add_paragraph()
    
    # Save template
    doc.save('templates/mccombs_template.docx')

def create_resume_template():
    document = Document()
    
    # Set page margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Define styles
    styles = document.styles
    
    # Name style
    name_style = styles.add_style('NameStyle', WD_STYLE_TYPE.PARAGRAPH)
    name_style.font.name = 'Arial'
    name_style.font.size = Pt(14)
    name_style.font.bold = True
    name_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_style.paragraph_format.space_after = Pt(2)
    
    # Contact info style
    contact_style = styles.add_style('ContactStyle', WD_STYLE_TYPE.PARAGRAPH)
    contact_style.font.name = 'Arial'
    contact_style.font.size = Pt(10)
    contact_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_style.paragraph_format.space_after = Pt(6)
    
    # Section heading style - made more prominent
    section_style = styles.add_style('SectionStyle', WD_STYLE_TYPE.PARAGRAPH)
    section_style.font.name = 'Arial'
    section_style.font.size = Pt(11)
    section_style.font.bold = True
    section_style.paragraph_format.space_before = Pt(4)
    section_style.paragraph_format.space_after = Pt(4)
    section_style.paragraph_format.border_bottom = True
    
    # Content style with better line spacing
    content_style = styles.add_style('ContentStyle', WD_STYLE_TYPE.PARAGRAPH)
    content_style.font.name = 'Arial'
    content_style.font.size = Pt(10)
    content_style.paragraph_format.space_after = Pt(3)
    content_style.paragraph_format.line_spacing = 1.05  # Slightly increased for readability
    
    # Add a style for education institution
    edu_style = styles.add_style('EducationStyle', WD_STYLE_TYPE.PARAGRAPH)
    edu_style.font.name = 'Arial'
    edu_style.font.size = Pt(10)
    edu_style.paragraph_format.space_after = Pt(3)
    edu_style.paragraph_format.line_spacing = 1.05
    
    return document

def format_resume(content, document):
    """
    Format the resume content with enhanced readability
    """
    try:
        previous_style = None
        for paragraph in content:
            text = paragraph.text.strip()
            if not text:  # Skip empty paragraphs
                continue
                
            p = document.add_paragraph()
            
            # Detect section headers
            if text.isupper() and '•' not in text and '@' not in text:
                p.style = 'SectionStyle'
                if previous_style:
                    p.paragraph_format.space_before = Pt(6)
                p.text = text
                previous_style = 'SectionStyle'
            
            # Format contact info
            elif '@' in text and '•' in text:
                p.style = 'ContactStyle'
                # Split contact info and rejoin with proper spacing
                contacts = [item.strip() for item in text.split('•')]
                p.text = ' • '.join(contacts)
                previous_style = 'ContactStyle'
            
            # Format name
            elif not document.paragraphs[:-1]:
                p.style = 'NameStyle'
                p.text = text.upper()
                previous_style = 'NameStyle'
            
            # Format regular content
            else:
                p.style = 'ContentStyle'
                
                # Education section special formatting
                if previous_style == 'SectionStyle' and 'University' in text:
                    p.style = 'EducationStyle'
                    parts = text.split('\t')
                    if len(parts) >= 2:
                        p.add_run(parts[0]).bold = True  # University name
                        p.add_run('\t' + parts[1])  # Degree
                        if len(parts) > 2:
                            p.add_run('\t' + parts[2])  # Date
                
                # Project or Experience entries
                elif text.startswith('•'):
                    # Remove bullet point for processing
                    text = text[1:].strip()
                    # Check for title - description format
                    if ' - ' in text:
                        title, description = text.split(' - ', 1)
                        p.add_run('• ').bold = True
                        p.add_run(title.strip()).bold = True
                        p.add_run(f" - {description.strip()}")
                    else:
                        p.add_run('• ').bold = True
                        p.add_run(text)
                
                # Course listings or other content
                else:
                    # Check if this is a course listing
                    if previous_style == 'SectionStyle' and ('Courses' in text or 'Skills' in text):
                        label, courses = text.split(':', 1) if ':' in text else (text, '')
                        p.add_run(f"{label}:").bold = True
                        p.add_run(courses)
                    else:
                        p.text = text
                
                previous_style = p.style.name

    except Exception as e:
        print(f"Error formatting resume: {str(e)}")
        raise

if __name__ == '__main__':
    setup_logging()
    ensure_directories()
    port = int(os.environ.get('PORT', 8080))
    if os.environ.get('FLASK_ENV') == 'production':
        app.run(host='0.0.0.0', port=port)
    else:
        app.run(debug=True, host='0.0.0.0', port=port)
